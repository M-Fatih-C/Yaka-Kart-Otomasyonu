#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Medar Yaka Kart Otomasyonu v3.0
==============================
Yeni Özellikler:
- Dosya listesi görünümü (sıralama, silme)
- Sürükle-bırak desteği
- İlerleme çubuğu
- Önizleme paneli
- Profil kaydetme/yükleme
- Gelişmiş çıktı ayarları
- Tema desteği (Açık/Koyu)
- İstatistik paneli
- Gelişmiş ayarlar (DPI, sayfa başı kart, boşluk)
"""

from pathlib import Path
from io import BytesIO
import threading
import os
import sys
import math
import zipfile
import shutil
import json
from datetime import datetime
from typing import List, Tuple, Optional, Dict, Any

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_SECTION
from PIL import Image, ImageTk

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import subprocess

# ================== GENEL AYARLAR ==================

if getattr(sys, "frozen", False):
    BASE_DIR = Path(sys.executable).resolve().parent
else:
    BASE_DIR = Path(__file__).resolve().parent

# Yapılandırma dosyaları
CONFIG_FILE = BASE_DIR / "config.json"
PROFILES_FILE = BASE_DIR / "profiles.json"
STATS_FILE = BASE_DIR / "stats.json"

# 7-Zip desteği
SEVEN_ZIP_SUPPORT = False
SEVEN_ZIP_PATH = None

def find_7zip():
    """7-Zip'i sistemde bul"""
    global SEVEN_ZIP_SUPPORT, SEVEN_ZIP_PATH
    
    possible_paths = [
        BASE_DIR / "7z.exe",
        Path("C:/Program Files/7-Zip/7z.exe"),
        Path("C:/Program Files (x86)/7-Zip/7z.exe"),
    ]
    
    for path in possible_paths:
        if path.exists():
            SEVEN_ZIP_PATH = str(path)
            SEVEN_ZIP_SUPPORT = True
            return True
    
    seven_zip_cmd = shutil.which("7z.exe")
    if seven_zip_cmd:
        SEVEN_ZIP_PATH = seven_zip_cmd
        SEVEN_ZIP_SUPPORT = True
        return True
    
    return False

find_7zip()

try:
    import rarfile
    RAR_SUPPORT = True
except ImportError:
    RAR_SUPPORT = False

# Sürükle-bırak desteği
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    DND_SUPPORT = True
except ImportError:
    DND_SUPPORT = False

ICON_PATH = BASE_DIR / "medar.ico"

# ================== VARSAYILAN AYARLAR ==================

DEFAULT_CONFIG = {
    "card_height_cm": 5.81,
    "card_width_cm": 9.2,
    "front_margins": {
        "top": 1.27,
        "bottom": 1.27,
        "left": 1.27,
        "right": 1.27
    },
    "back_margins": {
        "top": 1.27,
        "bottom": 1.27,
        "left": 0.7,
        "right": 1.27
    },
    "render_dpi": 300,
    "cards_per_page": 8,
    "card_spacing_cm": 0.0,
    "output_dir": str(BASE_DIR / "output"),
    "output_format": "docx",  # docx, pdf, both
    "filename_template": "kartlar_{date}_{time}",
    "theme": "light",
    "last_profile": "Varsayılan"
}

DEFAULT_PROFILES = {
    "Varsayılan": {
        "card_height_cm": 5.81,
        "card_width_cm": 9.2,
        "front_margins": {"top": 1.27, "bottom": 1.27, "left": 1.27, "right": 1.27},
        "back_margins": {"top": 1.27, "bottom": 1.27, "left": 0.7, "right": 1.27},
        "render_dpi": 300,
        "cards_per_page": 8
    },
    "Personel Kartı": {
        "card_height_cm": 5.5,
        "card_width_cm": 8.5,
        "front_margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
        "back_margins": {"top": 1.0, "bottom": 1.0, "left": 0.8, "right": 1.0},
        "render_dpi": 300,
        "cards_per_page": 8
    },
    "Ziyaretçi Kartı": {
        "card_height_cm": 5.0,
        "card_width_cm": 8.0,
        "front_margins": {"top": 1.5, "bottom": 1.5, "left": 1.5, "right": 1.5},
        "back_margins": {"top": 1.5, "bottom": 1.5, "left": 1.2, "right": 1.5},
        "render_dpi": 250,
        "cards_per_page": 8
    }
}

DEFAULT_STATS = {
    "total_cards": 0,
    "total_sessions": 0,
    "last_session_date": None,
    "last_session_cards": 0
}

# ================== TEMA TANIMLARI ==================

THEMES = {
    "light": {
        "bg": "#f5f5f5",
        "fg": "#333333",
        "frame_bg": "#ffffff",
        "accent": "#1976d2",
        "accent_hover": "#1565c0",
        "success": "#4CAF50",
        "warning": "#FF9800",
        "error": "#f44336",
        "button_fg": "#ffffff",
        "entry_bg": "#ffffff",
        "entry_fg": "#333333",
        "listbox_bg": "#ffffff",
        "listbox_fg": "#333333",
        "listbox_select_bg": "#1976d2",
        "listbox_select_fg": "#ffffff",
        "header_bg": "#1976d2",
        "header_fg": "#ffffff",
        "status_bg": "#f5f5f5",
        "tab_bg": "#e0e0e0",
        "tab_fg": "#333333"
    },
    "dark": {
        "bg": "#1e1e1e",
        "fg": "#e0e0e0",
        "frame_bg": "#2d2d2d",
        "accent": "#64b5f6",
        "accent_hover": "#42a5f5",
        "success": "#81c784",
        "warning": "#ffb74d",
        "error": "#e57373",
        "button_fg": "#1e1e1e",
        "entry_bg": "#3d3d3d",
        "entry_fg": "#e0e0e0",
        "listbox_bg": "#2d2d2d",
        "listbox_fg": "#e0e0e0",
        "listbox_select_bg": "#64b5f6",
        "listbox_select_fg": "#1e1e1e",
        "header_bg": "#37474f",
        "header_fg": "#ffffff",
        "status_bg": "#252525",
        "tab_bg": "#3d3d3d",
        "tab_fg": "#e0e0e0"
    }
}

# ================== YARDIMCI FONKSİYONLAR ==================

def load_config() -> Dict[str, Any]:
    """Yapılandırmayı yükle"""
    try:
        if CONFIG_FILE.exists():
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                config = json.load(f)
                # Eksik anahtarları varsayılanlarla doldur
                for key, value in DEFAULT_CONFIG.items():
                    if key not in config:
                        config[key] = value
                return config
    except Exception:
        pass
    return DEFAULT_CONFIG.copy()


def save_config(config: Dict[str, Any]):
    """Yapılandırmayı kaydet"""
    try:
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"Config kaydetme hatası: {e}")


def load_profiles() -> Dict[str, Dict]:
    """Profilleri yükle"""
    try:
        if PROFILES_FILE.exists():
            with open(PROFILES_FILE, 'r', encoding='utf-8') as f:
                profiles = json.load(f)
                # Varsayılan profilleri ekle
                for name, profile in DEFAULT_PROFILES.items():
                    if name not in profiles:
                        profiles[name] = profile
                return profiles
    except Exception:
        pass
    return DEFAULT_PROFILES.copy()


def save_profiles(profiles: Dict[str, Dict]):
    """Profilleri kaydet"""
    try:
        with open(PROFILES_FILE, 'w', encoding='utf-8') as f:
            json.dump(profiles, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"Profil kaydetme hatası: {e}")


def load_stats() -> Dict[str, Any]:
    """İstatistikleri yükle"""
    try:
        if STATS_FILE.exists():
            with open(STATS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception:
        pass
    return DEFAULT_STATS.copy()


def save_stats(stats: Dict[str, Any]):
    """İstatistikleri kaydet"""
    try:
        with open(STATS_FILE, 'w', encoding='utf-8') as f:
            json.dump(stats, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"İstatistik kaydetme hatası: {e}")


def update_stats(cards_created: int):
    """İstatistikleri güncelle"""
    stats = load_stats()
    stats["total_cards"] += cards_created
    stats["total_sessions"] += 1
    stats["last_session_date"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    stats["last_session_cards"] = cards_created
    save_stats(stats)
    return stats


# ================== ARŞİV ÇIKARMA FONKSİYONLARI ==================

TEMP_EXTRACT_DIR = BASE_DIR / "temp_extracted"


def extract_with_7zip(archive_path: Path, extract_to: Path):
    """7-Zip kullanarak arşivi çıkar"""
    if not SEVEN_ZIP_SUPPORT:
        return None
    
    try:
        cmd = [
            SEVEN_ZIP_PATH, 'x', str(archive_path),
            f'-o{extract_to}', '-y', '*.pdf',
        ]
        
        result = subprocess.run(
            cmd, capture_output=True, text=True,
            creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0
        )
        
        if result.returncode == 0:
            return list(extract_to.glob('**/*.pdf'))
        else:
            raise RuntimeError(f"7-Zip hata kodu: {result.returncode}")
    except Exception as e:
        raise RuntimeError(f"7-Zip ile çıkarma hatası: {str(e)}")


def extract_with_zipfile(archive_path: Path, extract_to: Path):
    """Python zipfile ile ZIP çıkar"""
    extracted_pdfs = []
    
    with zipfile.ZipFile(archive_path, 'r') as zip_ref:
        for file_info in zip_ref.namelist():
            if file_info.lower().endswith('.pdf'):
                safe_name = Path(file_info).name
                extract_path = extract_to / safe_name
                
                with zip_ref.open(file_info) as source:
                    with open(extract_path, 'wb') as target:
                        shutil.copyfileobj(source, target)
                
                extracted_pdfs.append(extract_path)
    
    return extracted_pdfs


def extract_with_rarfile(archive_path: Path, extract_to: Path):
    """rarfile ile RAR çıkar"""
    if not RAR_SUPPORT:
        raise RuntimeError("rarfile modülü yüklü değil")
    
    extracted_pdfs = []
    
    with rarfile.RarFile(archive_path, 'r') as rar_ref:
        for file_info in rar_ref.namelist():
            if file_info.lower().endswith('.pdf'):
                safe_name = Path(file_info).name
                extract_path = extract_to / safe_name
                
                with rar_ref.open(file_info) as source:
                    with open(extract_path, 'wb') as target:
                        shutil.copyfileobj(source, target)
                
                extracted_pdfs.append(extract_path)
    
    return extracted_pdfs


def extract_archive(archive_path: Path, extract_to: Path):
    """Arşivi çıkar"""
    archive_path = Path(archive_path)
    extract_to = Path(extract_to)
    extract_to.mkdir(parents=True, exist_ok=True)
    
    try:
        if SEVEN_ZIP_SUPPORT:
            extracted = extract_with_7zip(archive_path, extract_to)
            if extracted:
                return extracted
        
        if archive_path.suffix.lower() == '.zip':
            return extract_with_zipfile(archive_path, extract_to)
        
        elif archive_path.suffix.lower() == '.rar':
            if RAR_SUPPORT:
                return extract_with_rarfile(archive_path, extract_to)
            else:
                raise RuntimeError("RAR desteği yok! 7-Zip kurun.")
        
        else:
            raise ValueError(f"Desteklenmeyen format: {archive_path.suffix}")
    
    except Exception as e:
        raise RuntimeError(f"{archive_path.name} çıkarılırken hata: {str(e)}")


def clear_temp_directory():
    """Geçici klasörü temizle"""
    if TEMP_EXTRACT_DIR.exists():
        shutil.rmtree(TEMP_EXTRACT_DIR)
    TEMP_EXTRACT_DIR.mkdir(parents=True, exist_ok=True)


# ================== PDF İŞLEME FONKSİYONLARI ==================

def pdf_to_front_back(pdf_path: Path, dpi: int = 300) -> Tuple[Image.Image, Image.Image]:
    """PDF'den ön ve arka görüntüleri al"""
    doc = fitz.open(str(pdf_path))
    if len(doc) < 2:
        doc.close()
        raise ValueError(f"{pdf_path.name} içinde 2 sayfa yok.")

    zoom = dpi / 72
    mat = fitz.Matrix(zoom, zoom)

    def page_to_image(page):
        pix = page.get_pixmap(matrix=mat, alpha=False)
        return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    front = page_to_image(doc[0])
    back = page_to_image(doc[1])
    doc.close()
    return front, back


def get_pdf_preview(pdf_path: Path, max_size: Tuple[int, int] = (200, 150)) -> Optional[Image.Image]:
    """PDF'in önizleme görüntüsünü al"""
    try:
        doc = fitz.open(str(pdf_path))
        if len(doc) < 1:
            doc.close()
            return None
        
        page = doc[0]
        pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5), alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        doc.close()
        
        # Boyutlandır
        img.thumbnail(max_size, Image.Resampling.LANCZOS)
        return img
    except Exception:
        return None


def pil_to_stream(img: Image.Image) -> BytesIO:
    """PIL Image'ı BytesIO stream'e çevir"""
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def add_grid_page(doc: Document, images: List[Image.Image], rotate_degrees: int,
                  card_height_cm: float, card_width_cm: float,
                  cards_per_row: int = 2, reverse_rows: bool = False):
    """Görüntüleri grid halinde sayfaya ekle"""
    if not images:
        return

    max_cols = cards_per_row
    total = len(images)
    rows = math.ceil(total / max_cols)

    table = doc.add_table(rows=rows, cols=max_cols)
    table.autofit = False

    col_width = Cm(card_width_cm + 0.5)
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    idx = 0
    for r in range(rows):
        col_range = range(max_cols - 1, -1, -1) if reverse_rows else range(max_cols)
        
        for c in col_range:
            if idx >= total:
                break

            img = images[idx].rotate(rotate_degrees, expand=True)
            cell = table.rows[r].cells[c]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            stream = pil_to_stream(img)
            run.add_picture(stream, height=Cm(card_height_cm))
            idx += 1


def generate_doc_from_pdfs(pdf_paths: List[Path], card_height_cm: float, card_width_cm: float,
                           front_margins: Tuple, back_margins: Tuple,
                           render_dpi: int = 300, cards_per_page: int = 8,
                           output_path: Path = None,
                           progress_callback=None, status_callback=None) -> Path:
    """PDF'lerden Word dosyası oluştur"""
    
    if status_callback:
        status_callback("PDF'ler okunuyor...")

    output_dir = output_path.parent if output_path else BASE_DIR / "output"
    output_dir.mkdir(exist_ok=True, parents=True)
    
    if not output_path:
        output_path = output_dir / "kartlar.docx"

    # PDF'lerden görüntüleri al
    person_cards = []
    total_pdfs = len(pdf_paths)
    
    for i, pdf in enumerate(pdf_paths):
        try:
            front, back = pdf_to_front_back(pdf, dpi=render_dpi)
            person_cards.append((front, back))
            
            if progress_callback:
                progress_callback((i + 1) / total_pdfs * 50)  # İlk %50
            if status_callback:
                status_callback(f"Yüklendi: {pdf.name} ({i+1}/{total_pdfs})")
        except Exception as e:
            if status_callback:
                status_callback(f"HATA: {pdf.name} → {e}")

    if not person_cards:
        raise RuntimeError("Hiç geçerli PDF işlenemedi.")

    if status_callback:
        status_callback("Word dosyası oluşturuluyor...")

    doc = Document()

    # İlk section
    front_section = doc.sections[0]
    front_section.top_margin = Cm(front_margins[0])
    front_section.bottom_margin = Cm(front_margins[1])
    front_section.left_margin = Cm(front_margins[2])
    front_section.right_margin = Cm(front_margins[3])

    first_group = True
    total_groups = math.ceil(len(person_cards) / cards_per_page)

    for group_idx, i in enumerate(range(0, len(person_cards), cards_per_page)):
        group = person_cards[i:i + cards_per_page]

        # ÖN YÜZ
        if not first_group:
            front_section = doc.add_section(WD_SECTION.NEW_PAGE)
            front_section.top_margin = Cm(front_margins[0])
            front_section.bottom_margin = Cm(front_margins[1])
            front_section.left_margin = Cm(front_margins[2])
            front_section.right_margin = Cm(front_margins[3])

        front_images = [f for (f, _) in group]
        add_grid_page(doc, front_images, 90, card_height_cm, card_width_cm, reverse_rows=False)

        # ARKA YÜZ
        back_section = doc.add_section(WD_SECTION.NEW_PAGE)
        back_section.top_margin = Cm(back_margins[0])
        back_section.bottom_margin = Cm(back_margins[1])
        back_section.left_margin = Cm(back_margins[2])
        back_section.right_margin = Cm(back_margins[3])

        back_images = [b for (_, b) in group]
        add_grid_page(doc, back_images, 270, card_height_cm, card_width_cm, reverse_rows=True)

        first_group = False
        
        if progress_callback:
            progress_callback(50 + (group_idx + 1) / total_groups * 50)  # Son %50

    doc.save(output_path)
    return output_path


# ================== ANA UYGULAMA ==================

class YakaKartApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Medar Yaka Kart Otomasyonu v3.0")
        self.root.geometry("1100x900")
        self.root.minsize(900, 700)

        # Veri
        self.selected_files: List[Path] = []
        self.config = load_config()
        self.profiles = load_profiles()
        self.stats = load_stats()
        self.current_theme = self.config.get("theme", "light")
        self.preview_image = None

        # Tema uygula
        self.apply_theme()
        
        # Sürükle-bırak desteği
        if DND_SUPPORT:
            self.setup_dnd()

        # Ana layout
        self.create_ui()
        
        # Son profili yükle
        last_profile = self.config.get("last_profile", "Varsayılan")
        if last_profile in self.profiles:
            self.load_profile(last_profile)

    def apply_theme(self):
        """Temayı uygula"""
        theme = THEMES.get(self.current_theme, THEMES["light"])
        
        style = ttk.Style()
        style.theme_use('clam')
        
        # Notebook (tab) stilleri
        style.configure("TNotebook", background=theme["bg"])
        style.configure("TNotebook.Tab", 
                       background=theme["tab_bg"], 
                       foreground=theme["tab_fg"],
                       padding=[12, 6])
        style.map("TNotebook.Tab",
                 background=[("selected", theme["accent"])],
                 foreground=[("selected", theme["button_fg"])])
        
        # Progressbar stili
        style.configure("TProgressbar",
                       background=theme["accent"],
                       troughcolor=theme["entry_bg"])
        
        self.theme = theme
        self.root.configure(bg=theme["bg"])

    def setup_dnd(self):
        """Sürükle-bırak ayarla"""
        if DND_SUPPORT:
            self.root.drop_target_register(DND_FILES)
            self.root.dnd_bind('<<Drop>>', self.on_drop)

    def on_drop(self, event):
        """Dosya bırakıldığında"""
        files = self.root.tk.splitlist(event.data)
        pdf_files = []
        archive_files = []
        
        for f in files:
            path = Path(f)
            if path.suffix.lower() == '.pdf':
                pdf_files.append(path)
            elif path.suffix.lower() in ['.zip', '.rar', '.7z']:
                archive_files.append(path)
        
        if pdf_files:
            self.add_files_to_list(pdf_files)
        
        if archive_files:
            self.extract_and_add_archives(archive_files)

    def create_ui(self):
        """Ana arayüzü oluştur"""
        theme = self.theme

        # ========== ÜST PANEL (Header) ==========
        header_frame = tk.Frame(self.root, bg=theme["header_bg"])
        header_frame.pack(fill="x")

        # Sol: Başlık
        title_frame = tk.Frame(header_frame, bg=theme["header_bg"])
        title_frame.pack(side="left", padx=15, pady=10)

        tk.Label(
            title_frame,
            text="🖨️ Long-Edge Dupleks Baskı İçin Optimize Edilmiştir",
            bg=theme["header_bg"],
            fg=theme["header_fg"],
            font=("Arial", 11, "bold")
        ).pack(anchor="w")

        tk.Label(
            title_frame,
            text="📋 Yazıcı: Dupleks=Açık • Flip=Long Edge • Ölçek=%100 • Kağıt=A4",
            bg=theme["header_bg"],
            fg=theme["header_fg"],
            font=("Arial", 9)
        ).pack(anchor="w")

        # Sağ: Tema ve ayarlar
        settings_frame = tk.Frame(header_frame, bg=theme["header_bg"])
        settings_frame.pack(side="right", padx=15, pady=10)

        # Tema değiştirme butonu
        theme_icon = "🌙" if self.current_theme == "light" else "☀️"
        self.btn_theme = tk.Button(
            settings_frame,
            text=f"{theme_icon} Tema",
            command=self.toggle_theme,
            bg=theme["warning"],
            fg=theme["button_fg"],
            font=("Arial", 9, "bold"),
            relief="flat",
            cursor="hand2",
            padx=10,
            pady=5
        )
        self.btn_theme.pack(side="right", padx=5)

        # ========== ANA İÇERİK ==========
        main_frame = tk.Frame(self.root, bg=theme["bg"])
        main_frame.pack(fill="both", expand=True, padx=15, pady=10)

        # Sol panel (dosya listesi + önizleme)
        left_panel = tk.Frame(main_frame, bg=theme["bg"])
        left_panel.pack(side="left", fill="both", expand=True, padx=(0, 10))

        # Sağ panel (ayarlar)
        right_panel = tk.Frame(main_frame, bg=theme["bg"], width=350)
        right_panel.pack(side="right", fill="y")
        right_panel.pack_propagate(False)

        # ----- SOL PANEL İÇERİĞİ -----
        self.create_file_section(left_panel)
        self.create_preview_section(left_panel)
        self.create_log_section(left_panel)

        # ----- SAĞ PANEL İÇERİĞİ -----
        self.create_profile_section(right_panel)
        self.create_card_settings_section(right_panel)
        self.create_margin_section(right_panel)
        self.create_output_section(right_panel)
        self.create_stats_section(right_panel)

        # ========== ALT PANEL ==========
        self.create_bottom_panel()

    def create_file_section(self, parent):
        """Dosya seçimi bölümü"""
        theme = self.theme

        file_frame = tk.LabelFrame(
            parent,
            text=" 📁 Dosya Seçimi ",
            font=("Arial", 10, "bold"),
            bg=theme["frame_bg"],
            fg=theme["fg"],
            padx=10,
            pady=10
        )
        file_frame.pack(fill="x", pady=(0, 10))

        # Butonlar
        btn_container = tk.Frame(file_frame, bg=theme["frame_bg"])
        btn_container.pack(fill="x")

        self.btn_select_pdf = tk.Button(
            btn_container,
            text="📄 PDF Seç",
            command=self.select_pdfs,
            bg=theme["accent"],
            fg=theme["button_fg"],
            font=("Arial", 9, "bold"),
            relief="flat",
            cursor="hand2",
            padx=12,
            pady=6
        )
        self.btn_select_pdf.pack(side="left", padx=(0, 5))

        self.btn_select_archive = tk.Button(
            btn_container,
            text="📦 ZIP/RAR Aç",
            command=self.select_and_extract_archives,
            bg="#9C27B0",
            fg=theme["button_fg"],
            font=("Arial", 9, "bold"),
            relief="flat",
            cursor="hand2",
            padx=12,
            pady=6
        )
        self.btn_select_archive.pack(side="left", padx=(0, 5))

        self.btn_clear = tk.Button(
            btn_container,
            text="🗑️ Temizle",
            command=self.clear_file_list,
            bg=theme["error"],
            fg=theme["button_fg"],
            font=("Arial", 9, "bold"),
            relief="flat",
            cursor="hand2",
            padx=12,
            pady=6
        )
        self.btn_clear.pack(side="left")

        # Sürükle-bırak bilgisi
        if DND_SUPPORT:
            dnd_label = tk.Label(
                btn_container,
                text="📎 Sürükle-bırak destekleniyor",
                bg=theme["frame_bg"],
                fg=theme["success"],
                font=("Arial", 8, "italic")
            )
            dnd_label.pack(side="right")

        # Dosya listesi
        list_frame = tk.Frame(file_frame, bg=theme["frame_bg"])
        list_frame.pack(fill="both", expand=True, pady=(10, 0))

        # Scrollbar
        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side="right", fill="y")

        self.file_listbox = tk.Listbox(
            list_frame,
            height=8,
            font=("Consolas", 9),
            bg=theme["listbox_bg"],
            fg=theme["listbox_fg"],
            selectbackground=theme["listbox_select_bg"],
            selectforeground=theme["listbox_select_fg"],
            yscrollcommand=scrollbar.set,
            selectmode=tk.EXTENDED
        )
        self.file_listbox.pack(side="left", fill="both", expand=True)
        scrollbar.config(command=self.file_listbox.yview)

        # Listbox event'leri
        self.file_listbox.bind('<<ListboxSelect>>', self.on_file_select)
        self.file_listbox.bind('<Delete>', lambda e: self.remove_selected_files())

        # Liste kontrol butonları
        list_btn_frame = tk.Frame(file_frame, bg=theme["frame_bg"])
        list_btn_frame.pack(fill="x", pady=(5, 0))

        tk.Button(
            list_btn_frame,
            text="⬆️ Yukarı",
            command=self.move_file_up,
            bg=theme["tab_bg"],
            fg=theme["fg"],
            font=("Arial", 8),
            relief="flat",
            padx=8,
            pady=2
        ).pack(side="left", padx=(0, 5))

        tk.Button(
            list_btn_frame,
            text="⬇️ Aşağı",
            command=self.move_file_down,
            bg=theme["tab_bg"],
            fg=theme["fg"],
            font=("Arial", 8),
            relief="flat",
            padx=8,
            pady=2
        ).pack(side="left", padx=(0, 5))

        tk.Button(
            list_btn_frame,
            text="🗑️ Seçileni Sil",
            command=self.remove_selected_files,
            bg=theme["error"],
            fg=theme["button_fg"],
            font=("Arial", 8),
            relief="flat",
            padx=8,
            pady=2
        ).pack(side="left")

        self.lbl_file_count = tk.Label(
            list_btn_frame,
            text="0 dosya",
            bg=theme["frame_bg"],
            fg=theme["fg"],
            font=("Arial", 9, "bold")
        )
        self.lbl_file_count.pack(side="right")

    def create_preview_section(self, parent):
        """Önizleme bölümü"""
        theme = self.theme

        preview_frame = tk.LabelFrame(
            parent,
            text=" 👁️ Önizleme ",
            font=("Arial", 10, "bold"),
            bg=theme["frame_bg"],
            fg=theme["fg"],
            padx=10,
            pady=10
        )
        preview_frame.pack(fill="x", pady=(0, 10))

        # Önizleme canvas'ları
        preview_container = tk.Frame(preview_frame, bg=theme["frame_bg"])
        preview_container.pack(fill="x")

        # Ön yüz önizleme
        front_preview_frame = tk.Frame(preview_container, bg=theme["frame_bg"])
        front_preview_frame.pack(side="left", expand=True, fill="x", padx=(0, 10))

        tk.Label(
            front_preview_frame,
            text="Ön Yüz",
            bg=theme["frame_bg"],
            fg=theme["fg"],
            font=("Arial", 9, "bold")
        ).pack()

        self.front_canvas = tk.Canvas(
            front_preview_frame,
            width=180,
            height=120,
            bg=theme["entry_bg"],
            highlightthickness=1,
            highlightbackground=theme["tab_bg"]
        )
        self.front_canvas.pack(pady=5)

        # Arka yüz önizleme
        back_preview_frame = tk.Frame(preview_container, bg=theme["frame_bg"])
        back_preview_frame.pack(side="left", expand=True, fill="x")

        tk.Label(
            back_preview_frame,
            text="Arka Yüz",
            bg=theme["frame_bg"],
            fg=theme["fg"],
            font=("Arial", 9, "bold")
        ).pack()

        self.back_canvas = tk.Canvas(
            back_preview_frame,
            width=180,
            height=120,
            bg=theme["entry_bg"],
            highlightthickness=1,
            highlightbackground=theme["tab_bg"]
        )
        self.back_canvas.pack(pady=5)

        # Dosya adı label
        self.lbl_preview_name = tk.Label(
            preview_frame,
            text="Dosya seçin",
            bg=theme["frame_bg"],
            fg=theme["fg"],
            font=("Arial", 8, "italic")
        )
        self.lbl_preview_name.pack()

    def create_log_section(self, parent):
        """Log bölümü"""
        theme = self.theme

        log_frame = tk.LabelFrame(
            parent,
            text=" 📜 İşlem Logu ",
            font=("Arial", 10, "bold"),
            bg=theme["frame_bg"],
            fg=theme["fg"],
            padx=10,
            pady=5
        )
        log_frame.pack(fill="both", expand=True)

        log_scroll = tk.Scrollbar(log_frame)
        log_scroll.pack(side="right", fill="y")

        self.log_list = tk.Listbox(
            log_frame,
            height=6,
            font=("Consolas", 9),
            bg=theme["listbox_bg"],
            fg=theme["listbox_fg"],
            yscrollcommand=log_scroll.set
        )
        self.log_list.pack(fill="both", expand=True)
        log_scroll.config(command=self.log_list.yview)

    def create_profile_section(self, parent):
        """Profil bölümü"""
        theme = self.theme

        profile_frame = tk.LabelFrame(
            parent,
            text=" 📋 Profiller ",
            font=("Arial", 10, "bold"),
            bg=theme["frame_bg"],
            fg=theme["fg"],
            padx=10,
            pady=10
        )
        profile_frame.pack(fill="x", pady=(0, 10))

        # Profil seçimi
        select_frame = tk.Frame(profile_frame, bg=theme["frame_bg"])
        select_frame.pack(fill="x")

        tk.Label(
            select_frame,
            text="Profil:",
            bg=theme["frame_bg"],
            fg=theme["fg"],
            font=("Arial", 9)
        ).pack(side="left")

        self.profile_var = tk.StringVar(value=self.config.get("last_profile", "Varsayılan"))
        self.profile_combo = ttk.Combobox(
            select_frame,
            textvariable=self.profile_var,
            values=list(self.profiles.keys()),
            state="readonly",
            width=15
        )
        self.profile_combo.pack(side="left", padx=5)
        self.profile_combo.bind('<<ComboboxSelected>>', self.on_profile_change)

        # Profil butonları
        btn_frame = tk.Frame(profile_frame, bg=theme["frame_bg"])
        btn_frame.pack(fill="x", pady=(10, 0))

        tk.Button(
            btn_frame,
            text="💾 Kaydet",
            command=self.save_current_as_profile,
            bg=theme["success"],
            fg=theme["button_fg"],
            font=("Arial", 8, "bold"),
            relief="flat",
            padx=8,
            pady=3
        ).pack(side="left", padx=(0, 5))

        tk.Button(
            btn_frame,
            text="🗑️ Sil",
            command=self.delete_profile,
            bg=theme["error"],
            fg=theme["button_fg"],
            font=("Arial", 8, "bold"),
            relief="flat",
            padx=8,
            pady=3
        ).pack(side="left")

    def create_card_settings_section(self, parent):
        """Kart ayarları bölümü"""
        theme = self.theme

        card_frame = tk.LabelFrame(
            parent,
            text=" 📐 Kart Ayarları ",
            font=("Arial", 10, "bold"),
            bg=theme["frame_bg"],
            fg=theme["fg"],
            padx=10,
            pady=10
        )
        card_frame.pack(fill="x", pady=(0, 10))

        # Boyutlar
        size_frame = tk.Frame(card_frame, bg=theme["frame_bg"])
        size_frame.pack(fill="x")

        tk.Label(size_frame, text="Yükseklik (cm):", bg=theme["frame_bg"], 
                fg=theme["fg"], font=("Arial", 9)).grid(row=0, column=0, sticky="w", pady=2)
        self.entry_height = tk.Entry(size_frame, width=10, font=("Arial", 9),
                                     bg=theme["entry_bg"], fg=theme["entry_fg"])
        self.entry_height.grid(row=0, column=1, padx=5, pady=2)
        self.entry_height.insert(0, str(self.config["card_height_cm"]))

        tk.Label(size_frame, text="Genişlik (cm):", bg=theme["frame_bg"],
                fg=theme["fg"], font=("Arial", 9)).grid(row=1, column=0, sticky="w", pady=2)
        self.entry_width = tk.Entry(size_frame, width=10, font=("Arial", 9),
                                    bg=theme["entry_bg"], fg=theme["entry_fg"])
        self.entry_width.grid(row=1, column=1, padx=5, pady=2)
        self.entry_width.insert(0, str(self.config["card_width_cm"]))

        # Gelişmiş ayarlar
        advanced_frame = tk.Frame(card_frame, bg=theme["frame_bg"])
        advanced_frame.pack(fill="x", pady=(10, 0))

        tk.Label(advanced_frame, text="DPI Kalitesi:", bg=theme["frame_bg"],
                fg=theme["fg"], font=("Arial", 9)).grid(row=0, column=0, sticky="w", pady=2)
        self.dpi_var = tk.StringVar(value=str(self.config.get("render_dpi", 300)))
        dpi_combo = ttk.Combobox(
            advanced_frame,
            textvariable=self.dpi_var,
            values=["150", "200", "250", "300", "350", "400"],
            width=8,
            state="readonly"
        )
        dpi_combo.grid(row=0, column=1, padx=5, pady=2)

        tk.Label(advanced_frame, text="Sayfa başı kart:", bg=theme["frame_bg"],
                fg=theme["fg"], font=("Arial", 9)).grid(row=1, column=0, sticky="w", pady=2)
        self.cards_per_page_var = tk.StringVar(value=str(self.config.get("cards_per_page", 8)))
        cards_combo = ttk.Combobox(
            advanced_frame,
            textvariable=self.cards_per_page_var,
            values=["4", "6", "8"],
            width=8,
            state="readonly"
        )
        cards_combo.grid(row=1, column=1, padx=5, pady=2)

    def create_margin_section(self, parent):
        """Kenar boşlukları bölümü"""
        theme = self.theme

        margin_frame = tk.LabelFrame(
            parent,
            text=" 📄 Kenar Boşlukları (cm) ",
            font=("Arial", 10, "bold"),
            bg=theme["frame_bg"],
            fg=theme["fg"],
            padx=10,
            pady=5
        )
        margin_frame.pack(fill="x", pady=(0, 10))

        notebook = ttk.Notebook(margin_frame)
        notebook.pack(fill="x")

        # Ön yüz tab
        front_tab = tk.Frame(notebook, bg=theme["frame_bg"])
        notebook.add(front_tab, text=" Ön Yüz ")
        self.front_margin_entries = self.create_margin_inputs(
            front_tab, self.config["front_margins"]
        )

        # Arka yüz tab
        back_tab = tk.Frame(notebook, bg=theme["frame_bg"])
        notebook.add(back_tab, text=" Arka Yüz ")
        self.back_margin_entries = self.create_margin_inputs(
            back_tab, self.config["back_margins"]
        )

    def create_margin_inputs(self, parent, values: Dict) -> Dict[str, tk.Entry]:
        """Kenar boşluğu input'ları oluştur"""
        theme = self.theme
        entries = {}

        container = tk.Frame(parent, bg=theme["frame_bg"])
        container.pack(padx=10, pady=10)

        labels = [("Üst:", "top"), ("Alt:", "bottom"), ("Sol:", "left"), ("Sağ:", "right")]
        
        for i, (label, key) in enumerate(labels):
            row, col = i // 2, (i % 2) * 2
            tk.Label(container, text=label, bg=theme["frame_bg"], fg=theme["fg"],
                    font=("Arial", 9)).grid(row=row, column=col, sticky="e", padx=2, pady=3)
            entries[key] = tk.Entry(container, width=8, font=("Arial", 9),
                                   bg=theme["entry_bg"], fg=theme["entry_fg"])
            entries[key].grid(row=row, column=col+1, padx=2, pady=3)
            entries[key].insert(0, str(values.get(key, 1.27)))

        return entries

    def create_output_section(self, parent):
        """Çıktı ayarları bölümü"""
        theme = self.theme

        output_frame = tk.LabelFrame(
            parent,
            text=" 📁 Çıktı Ayarları ",
            font=("Arial", 10, "bold"),
            bg=theme["frame_bg"],
            fg=theme["fg"],
            padx=10,
            pady=10
        )
        output_frame.pack(fill="x", pady=(0, 10))

        # Çıktı klasörü
        dir_frame = tk.Frame(output_frame, bg=theme["frame_bg"])
        dir_frame.pack(fill="x")

        tk.Label(dir_frame, text="Klasör:", bg=theme["frame_bg"],
                fg=theme["fg"], font=("Arial", 9)).pack(side="left")

        self.output_dir_var = tk.StringVar(value=self.config.get("output_dir", str(BASE_DIR / "output")))
        self.entry_output_dir = tk.Entry(
            dir_frame, textvariable=self.output_dir_var,
            font=("Arial", 8), width=20,
            bg=theme["entry_bg"], fg=theme["entry_fg"]
        )
        self.entry_output_dir.pack(side="left", padx=5, fill="x", expand=True)

        tk.Button(
            dir_frame,
            text="📂",
            command=self.select_output_dir,
            bg=theme["tab_bg"],
            fg=theme["fg"],
            font=("Arial", 9),
            relief="flat",
            padx=5
        ).pack(side="left")

        # Dosya adı şablonu
        name_frame = tk.Frame(output_frame, bg=theme["frame_bg"])
        name_frame.pack(fill="x", pady=(5, 0))

        tk.Label(name_frame, text="Dosya adı:", bg=theme["frame_bg"],
                fg=theme["fg"], font=("Arial", 9)).pack(side="left")

        self.filename_var = tk.StringVar(value=self.config.get("filename_template", "kartlar_{date}_{time}"))
        self.entry_filename = tk.Entry(
            name_frame, textvariable=self.filename_var,
            font=("Arial", 8), width=25,
            bg=theme["entry_bg"], fg=theme["entry_fg"]
        )
        self.entry_filename.pack(side="left", padx=5)

        tk.Label(
            name_frame,
            text="{date}, {time}",
            bg=theme["frame_bg"],
            fg=theme["fg"],
            font=("Arial", 7, "italic")
        ).pack(side="left")

    def create_stats_section(self, parent):
        """İstatistik bölümü"""
        theme = self.theme

        stats_frame = tk.LabelFrame(
            parent,
            text=" 📊 İstatistikler ",
            font=("Arial", 10, "bold"),
            bg=theme["frame_bg"],
            fg=theme["fg"],
            padx=10,
            pady=10
        )
        stats_frame.pack(fill="x")

        self.lbl_stats = tk.Label(
            stats_frame,
            text=self.get_stats_text(),
            bg=theme["frame_bg"],
            fg=theme["fg"],
            font=("Arial", 9),
            justify="left"
        )
        self.lbl_stats.pack(anchor="w")

    def create_bottom_panel(self):
        """Alt panel (ilerleme çubuğu, butonlar, durum)"""
        theme = self.theme

        bottom_frame = tk.Frame(self.root, bg=theme["bg"])
        bottom_frame.pack(fill="x", padx=15, pady=(0, 10))

        # İlerleme çubuğu
        progress_frame = tk.Frame(bottom_frame, bg=theme["bg"])
        progress_frame.pack(fill="x", pady=(0, 10))

        self.progress_var = tk.DoubleVar(value=0)
        self.progress_bar = ttk.Progressbar(
            progress_frame,
            variable=self.progress_var,
            maximum=100,
            length=400
        )
        self.progress_bar.pack(side="left", fill="x", expand=True, padx=(0, 10))

        self.lbl_progress = tk.Label(
            progress_frame,
            text="0%",
            bg=theme["bg"],
            fg=theme["fg"],
            font=("Arial", 9, "bold"),
            width=5
        )
        self.lbl_progress.pack(side="left")

        # Ana buton
        self.btn_run = tk.Button(
            bottom_frame,
            text="🚀 Kimlikleri Oluştur",
            command=self.run_generation,
            bg=theme["success"],
            fg=theme["button_fg"],
            font=("Arial", 12, "bold"),
            relief="flat",
            cursor="hand2",
            padx=40,
            pady=12
        )
        self.btn_run.pack(pady=(0, 10))

        # Durum çubuğu
        status_frame = tk.Frame(self.root, bg=theme["status_bg"], relief="sunken", borderwidth=1)
        status_frame.pack(fill="x", side="bottom")

        self.lbl_status = tk.Label(
            status_frame,
            text="⚡ Durum: Hazır",
            bg=theme["status_bg"],
            fg=theme["fg"],
            font=("Arial", 9),
            anchor="w"
        )
        self.lbl_status.pack(padx=10, pady=5, fill="x")

    # ========== DOSYA İŞLEMLERİ ==========

    def select_pdfs(self):
        """PDF dosyalarını seç"""
        files = filedialog.askopenfilenames(
            title="PDF Dosyalarını Seç",
            filetypes=[("PDF files", "*.pdf")]
        )
        if files:
            self.add_files_to_list([Path(f) for f in files])

    def select_and_extract_archives(self):
        """Arşiv dosyalarını seç ve çıkar"""
        if not SEVEN_ZIP_SUPPORT and not RAR_SUPPORT:
            messagebox.showwarning(
                "Uyarı",
                "Arşiv desteği için 7-Zip kurulu olmalı.\n"
                "https://www.7-zip.org/download.html"
            )

        filetypes = [
            ("Tüm Arşivler", "*.zip *.rar *.7z"),
            ("ZIP files", "*.zip"),
            ("RAR files", "*.rar"),
            ("7Z files", "*.7z")
        ]

        files = filedialog.askopenfilenames(
            title="Arşiv Dosyalarını Seç",
            filetypes=filetypes
        )

        if files:
            self.extract_and_add_archives([Path(f) for f in files])

    def extract_and_add_archives(self, archive_files: List[Path]):
        """Arşivleri çıkar ve ekle"""
        self.add_log("📦 Arşiv çıkarma başladı...")
        self.set_status("Arşivler çıkarılıyor...")
        self.disable_buttons()

        def worker():
            try:
                clear_temp_directory()
                all_extracted = []

                for archive_path in archive_files:
                    self.thread_safe_log(f"📂 Açılıyor: {archive_path.name}")
                    try:
                        extracted = extract_archive(archive_path, TEMP_EXTRACT_DIR)
                        all_extracted.extend(extracted)
                        self.thread_safe_log(f"  ✓ {len(extracted)} PDF çıkarıldı")
                    except Exception as e:
                        self.thread_safe_log(f"  ✗ HATA: {str(e)}")

                if all_extracted:
                    self.root.after(0, lambda: self.add_files_to_list(all_extracted))
                    self.thread_safe_log(f"✅ Toplam {len(all_extracted)} PDF eklendi")
                    self.thread_safe_status("Arşivler çıkarıldı")
                else:
                    self.thread_safe_log("⚠️ Hiç PDF bulunamadı")
                    self.thread_safe_status("Hazır")

            except Exception as e:
                self.thread_safe_log(f"❌ Hata: {str(e)}")
                self.thread_safe_status("Hata oluştu")
            finally:
                self.root.after(0, self.enable_buttons)

        threading.Thread(target=worker, daemon=True).start()

    def add_files_to_list(self, files: List[Path]):
        """Dosyaları listeye ekle"""
        for f in files:
            if f not in self.selected_files:
                self.selected_files.append(f)
                self.file_listbox.insert(tk.END, f.name)
                self.add_log(f"✓ Eklendi: {f.name}")

        self.update_file_count()

    def remove_selected_files(self):
        """Seçili dosyaları sil"""
        selected = list(self.file_listbox.curselection())
        if not selected:
            return

        for idx in reversed(selected):
            self.file_listbox.delete(idx)
            del self.selected_files[idx]

        self.update_file_count()
        self.clear_preview()

    def clear_file_list(self):
        """Tüm dosyaları temizle"""
        if not self.selected_files:
            return

        if messagebox.askyesno("Onay", f"{len(self.selected_files)} dosya silinecek. Emin misiniz?"):
            self.selected_files.clear()
            self.file_listbox.delete(0, tk.END)
            self.update_file_count()
            self.clear_preview()
            self.add_log("🗑️ Liste temizlendi")

    def move_file_up(self):
        """Seçili dosyayı yukarı taşı"""
        selected = self.file_listbox.curselection()
        if not selected or selected[0] == 0:
            return

        idx = selected[0]
        self.selected_files[idx], self.selected_files[idx-1] = \
            self.selected_files[idx-1], self.selected_files[idx]

        text = self.file_listbox.get(idx)
        self.file_listbox.delete(idx)
        self.file_listbox.insert(idx-1, text)
        self.file_listbox.selection_set(idx-1)

    def move_file_down(self):
        """Seçili dosyayı aşağı taşı"""
        selected = self.file_listbox.curselection()
        if not selected or selected[0] >= len(self.selected_files) - 1:
            return

        idx = selected[0]
        self.selected_files[idx], self.selected_files[idx+1] = \
            self.selected_files[idx+1], self.selected_files[idx]

        text = self.file_listbox.get(idx)
        self.file_listbox.delete(idx)
        self.file_listbox.insert(idx+1, text)
        self.file_listbox.selection_set(idx+1)

    def update_file_count(self):
        """Dosya sayısını güncelle"""
        count = len(self.selected_files)
        self.lbl_file_count.config(
            text=f"{count} dosya",
            fg=self.theme["success"] if count > 0 else self.theme["fg"]
        )

    def on_file_select(self, event):
        """Dosya seçildiğinde önizleme göster"""
        selected = self.file_listbox.curselection()
        if not selected:
            return

        idx = selected[0]
        pdf_path = self.selected_files[idx]
        self.show_preview(pdf_path)

    def show_preview(self, pdf_path: Path):
        """PDF önizlemesi göster"""
        try:
            doc = fitz.open(str(pdf_path))
            if len(doc) < 2:
                doc.close()
                return

            # Ön yüz
            front_pix = doc[0].get_pixmap(matrix=fitz.Matrix(0.3, 0.3), alpha=False)
            front_img = Image.frombytes("RGB", [front_pix.width, front_pix.height], front_pix.samples)
            front_img.thumbnail((180, 120), Image.Resampling.LANCZOS)

            # Arka yüz
            back_pix = doc[1].get_pixmap(matrix=fitz.Matrix(0.3, 0.3), alpha=False)
            back_img = Image.frombytes("RGB", [back_pix.width, back_pix.height], back_pix.samples)
            back_img.thumbnail((180, 120), Image.Resampling.LANCZOS)

            doc.close()

            # Canvas'a çiz
            self.front_photo = ImageTk.PhotoImage(front_img)
            self.back_photo = ImageTk.PhotoImage(back_img)

            self.front_canvas.delete("all")
            self.back_canvas.delete("all")

            self.front_canvas.create_image(90, 60, image=self.front_photo)
            self.back_canvas.create_image(90, 60, image=self.back_photo)

            self.lbl_preview_name.config(text=pdf_path.name)

        except Exception as e:
            self.add_log(f"Önizleme hatası: {e}")

    def clear_preview(self):
        """Önizlemeyi temizle"""
        self.front_canvas.delete("all")
        self.back_canvas.delete("all")
        self.lbl_preview_name.config(text="Dosya seçin")

    # ========== PROFİL İŞLEMLERİ ==========

    def on_profile_change(self, event=None):
        """Profil değiştiğinde"""
        profile_name = self.profile_var.get()
        self.load_profile(profile_name)

    def load_profile(self, profile_name: str):
        """Profili yükle"""
        if profile_name not in self.profiles:
            return

        profile = self.profiles[profile_name]

        # Kart boyutları
        self.entry_height.delete(0, tk.END)
        self.entry_height.insert(0, str(profile.get("card_height_cm", 5.81)))

        self.entry_width.delete(0, tk.END)
        self.entry_width.insert(0, str(profile.get("card_width_cm", 9.2)))

        # DPI ve sayfa başı kart
        self.dpi_var.set(str(profile.get("render_dpi", 300)))
        self.cards_per_page_var.set(str(profile.get("cards_per_page", 8)))

        # Kenar boşlukları
        front_margins = profile.get("front_margins", {})
        back_margins = profile.get("back_margins", {})

        for key, entry in self.front_margin_entries.items():
            entry.delete(0, tk.END)
            entry.insert(0, str(front_margins.get(key, 1.27)))

        for key, entry in self.back_margin_entries.items():
            entry.delete(0, tk.END)
            entry.insert(0, str(back_margins.get(key, 1.27)))

        self.config["last_profile"] = profile_name
        save_config(self.config)
        self.add_log(f"📋 Profil yüklendi: {profile_name}")

    def save_current_as_profile(self):
        """Mevcut ayarları profil olarak kaydet"""
        name = tk.simpledialog.askstring(
            "Profil Kaydet",
            "Profil adı:",
            initialvalue=self.profile_var.get()
        )

        if not name:
            return

        profile = {
            "card_height_cm": float(self.entry_height.get().replace(",", ".")),
            "card_width_cm": float(self.entry_width.get().replace(",", ".")),
            "render_dpi": int(self.dpi_var.get()),
            "cards_per_page": int(self.cards_per_page_var.get()),
            "front_margins": {
                key: float(entry.get().replace(",", "."))
                for key, entry in self.front_margin_entries.items()
            },
            "back_margins": {
                key: float(entry.get().replace(",", "."))
                for key, entry in self.back_margin_entries.items()
            }
        }

        self.profiles[name] = profile
        save_profiles(self.profiles)

        # Combo'yu güncelle
        self.profile_combo['values'] = list(self.profiles.keys())
        self.profile_var.set(name)

        self.add_log(f"💾 Profil kaydedildi: {name}")
        messagebox.showinfo("Başarılı", f"'{name}' profili kaydedildi.")

    def delete_profile(self):
        """Profili sil"""
        name = self.profile_var.get()

        if name in ["Varsayılan", "Personel Kartı", "Ziyaretçi Kartı"]:
            messagebox.showwarning("Uyarı", "Varsayılan profiller silinemez.")
            return

        if messagebox.askyesno("Onay", f"'{name}' profili silinecek. Emin misiniz?"):
            del self.profiles[name]
            save_profiles(self.profiles)

            self.profile_combo['values'] = list(self.profiles.keys())
            self.profile_var.set("Varsayılan")
            self.load_profile("Varsayılan")

            self.add_log(f"🗑️ Profil silindi: {name}")

    # ========== ÇIKTI AYARLARI ==========

    def select_output_dir(self):
        """Çıktı klasörü seç"""
        dir_path = filedialog.askdirectory(
            title="Çıktı Klasörü Seç",
            initialdir=self.output_dir_var.get()
        )
        if dir_path:
            self.output_dir_var.set(dir_path)

    def get_output_path(self) -> Path:
        """Çıktı dosya yolunu oluştur"""
        output_dir = Path(self.output_dir_var.get())
        output_dir.mkdir(parents=True, exist_ok=True)

        template = self.filename_var.get()
        now = datetime.now()

        filename = template.replace("{date}", now.strftime("%Y%m%d"))
        filename = filename.replace("{time}", now.strftime("%H%M%S"))
        filename = f"{filename}.docx"

        return output_dir / filename

    # ========== İSTATİSTİKLER ==========

    def get_stats_text(self) -> str:
        """İstatistik metnini oluştur"""
        stats = self.stats
        text = f"📈 Toplam: {stats['total_cards']} kart\n"
        text += f"📊 Oturum: {stats['total_sessions']} işlem\n"

        if stats['last_session_date']:
            text += f"🕐 Son: {stats['last_session_date']}\n"
            text += f"   ({stats['last_session_cards']} kart)"

        return text

    def update_stats_display(self):
        """İstatistik görünümünü güncelle"""
        self.stats = load_stats()
        self.lbl_stats.config(text=self.get_stats_text())

    # ========== TEMA ==========

    def toggle_theme(self):
        """Temayı değiştir"""
        self.current_theme = "dark" if self.current_theme == "light" else "light"
        self.config["theme"] = self.current_theme
        save_config(self.config)

        # Uygulamayı yeniden başlat uyarısı
        messagebox.showinfo(
            "Tema Değiştirildi",
            "Tema değişikliği için uygulama yeniden başlatılmalı."
        )

    # ========== YARDIMCI METODLAR ==========

    def add_log(self, message: str):
        """Log ekle"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_list.insert(tk.END, f"[{timestamp}] {message}")
        self.log_list.see(tk.END)

    def set_status(self, text: str):
        """Durum güncelle"""
        self.lbl_status.config(text=f"⚡ Durum: {text}")
        self.root.update_idletasks()

    def set_progress(self, value: float):
        """İlerleme güncelle"""
        self.progress_var.set(value)
        self.lbl_progress.config(text=f"{int(value)}%")
        self.root.update_idletasks()

    def disable_buttons(self):
        """Butonları devre dışı bırak"""
        self.btn_select_pdf.config(state="disabled")
        self.btn_select_archive.config(state="disabled")
        self.btn_clear.config(state="disabled")
        self.btn_run.config(state="disabled")

    def enable_buttons(self):
        """Butonları etkinleştir"""
        self.btn_select_pdf.config(state="normal")
        self.btn_select_archive.config(state="normal")
        self.btn_clear.config(state="normal")
        self.btn_run.config(state="normal")

    # Thread-safe metodlar
    def thread_safe_log(self, text: str):
        self.root.after(0, self.add_log, text)

    def thread_safe_status(self, text: str):
        self.root.after(0, self.set_status, text)

    def thread_safe_progress(self, value: float):
        self.root.after(0, self.set_progress, value)

    # ========== ANA İŞLEM ==========

    def get_margin_values(self, entries: Dict[str, tk.Entry]) -> Tuple:
        """Kenar boşluklarını al"""
        try:
            top = float(entries['top'].get().replace(",", "."))
            bottom = float(entries['bottom'].get().replace(",", "."))
            left = float(entries['left'].get().replace(",", "."))
            right = float(entries['right'].get().replace(",", "."))

            if any(v < 0 for v in [top, bottom, left, right]):
                raise ValueError("Negatif değer")

            return (top, bottom, left, right)
        except ValueError:
            raise ValueError("Kenar boşlukları pozitif sayı olmalıdır")

    def run_generation(self):
        """Kimlik oluştur"""
        if not self.selected_files:
            messagebox.showwarning("Uyarı", "Lütfen önce PDF dosyalarını seçin.")
            return

        # Değerleri oku
        try:
            h = float(self.entry_height.get().replace(",", "."))
            w = float(self.entry_width.get().replace(",", "."))
            if h <= 0 or w <= 0:
                raise ValueError()
        except ValueError:
            messagebox.showerror("Hata", "Kart boyutları pozitif sayı olmalıdır.")
            return

        try:
            front_margins = self.get_margin_values(self.front_margin_entries)
            back_margins = self.get_margin_values(self.back_margin_entries)
        except ValueError as e:
            messagebox.showerror("Hata", str(e))
            return

        dpi = int(self.dpi_var.get())
        cards_per_page = int(self.cards_per_page_var.get())
        output_path = self.get_output_path()

        self.disable_buttons()
        self.set_status("İşleniyor...")
        self.set_progress(0)

        def worker():
            try:
                generate_doc_from_pdfs(
                    self.selected_files,
                    card_height_cm=h,
                    card_width_cm=w,
                    front_margins=front_margins,
                    back_margins=back_margins,
                    render_dpi=dpi,
                    cards_per_page=cards_per_page,
                    output_path=output_path,
                    progress_callback=self.thread_safe_progress,
                    status_callback=self.thread_safe_status
                )

                # İstatistikleri güncelle
                cards_created = len(self.selected_files)
                update_stats(cards_created)
                self.root.after(0, self.update_stats_display)

                self.thread_safe_progress(100)
                self.thread_safe_status("Tamamlandı! ✓")

                self.root.after(0, lambda: messagebox.showinfo(
                    "✅ İşlem Tamamlandı",
                    f"Kimlikler başarıyla oluşturuldu!\n\n"
                    f"📁 Dosya: {output_path}\n\n"
                    f"🖨️ Yazıcı Ayarları:\n"
                    f"  ✓ Dupleks: Long Edge\n"
                    f"  ✓ Ölçek: %100\n"
                    f"  ✓ Kağıt: A4\n\n"
                    f"📋 {cards_created} kart oluşturuldu"
                ))

                # Klasörü aç
                try:
                    os.startfile(output_path.parent)
                except:
                    pass

            except Exception as e:
                self.thread_safe_status("Hata oluştu!")
                self.root.after(0, lambda: messagebox.showerror("❌ Hata", str(e)))
            finally:
                self.root.after(0, self.enable_buttons)

        threading.Thread(target=worker, daemon=True).start()


# ================== BAŞLATMA ==================

def main():
    # Sürükle-bırak desteği varsa TkinterDnD kullan
    if DND_SUPPORT:
        root = TkinterDnD.Tk()
    else:
        root = tk.Tk()

    # simpledialog import (profil kaydetme için)
    import tkinter.simpledialog as simpledialog
    tk.simpledialog = simpledialog

    try:
        root.iconbitmap(str(ICON_PATH))
    except:
        pass

    # Başlangıç mesajları
    print("=" * 50)
    print("Medar Yaka Kart Otomasyonu v3.0")
    print("=" * 50)

    if SEVEN_ZIP_SUPPORT:
        print(f"✅ 7-Zip: {SEVEN_ZIP_PATH}")
    else:
        print("⚠️ 7-Zip bulunamadı")

    if DND_SUPPORT:
        print("✅ Sürükle-bırak desteği aktif")
    else:
        print("⚠️ Sürükle-bırak için: pip install tkinterdnd2")

    print("=" * 50)

    app = YakaKartApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()